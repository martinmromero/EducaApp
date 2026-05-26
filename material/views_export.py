"""
views_export.py
Exportación de exámenes a DOCX y PDF.

Librerías utilizadas:
  - python-docx  : genera documentos Word (Office Open XML) en memoria.
  - xhtml2pdf    : convierte HTML/CSS a PDF vía ReportLab, sin dependencias
                   binarias del sistema (compatible con Render PaaS).

Ambas respuestas se arman en un io.BytesIO y se sirven como descarga HTTP;
ningún archivo se escribe a disco.
"""

import io
import base64

from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string

from .models import Exam


# ──────────────────────────────────────────────────────────────────────────────
# Helpers internos
# ──────────────────────────────────────────────────────────────────────────────

_TIPO_EXAMEN_LABELS = {
    'final': 'Final',
    'parcial': 'Parcial',
    '1er_parcial': '1er Parcial',
    '2do_parcial': '2do Parcial',
    '3er_parcial': '3er Parcial',
    'recuperatorio': 'Recuperatorio',
    'practico': 'Práctico',
}

_TIPO_MODALIDAD_LABELS = {'individual': 'Individual', 'grupal': 'Grupal'}

_OPTION_LETTERS = 'abcdefghijklmnopqrstuvwxyz'


def _build_export_context(examen, con_respuestas=False):
    """
    Construye el contexto de exportación con exactamente las mismas claves que
    ver_examen() pasa a su template, más la variable `con_respuestas`.
    Esto garantiza que PDF, DOCX y la vista browser usen la misma estructura.
    """
    questions_texts = []
    for q in examen.questions.order_by('pk').all():
        questions_texts.append({
            'text': q.question_text,
            'type': q.question_type,
            'options': q.options or [],
            'question_image_b64': q.question_image_b64 or '',
            'answer_text': q.answer_text or '',
            'answer_image_b64': q.answer_image_b64 or '',
        })

    outcomes_texts = [o.description for o in examen.learning_outcomes.all()]
    if not outcomes_texts and examen.outcomes_snapshot:
        outcomes_texts = list(examen.outcomes_snapshot)

    topics_texts = [t.name for t in examen.topics.all()]
    if not topics_texts and examen.topics_snapshot:
        topics_texts = list(examen.topics_snapshot)

    if examen.professor:
        professor_name = examen.professor.get_full_name() or examen.professor.username
    else:
        professor_name = '-'

    modalidad_list = [m.strip() for m in (examen.resolution_time or '').split(',') if m.strip()]

    return {
        # ── mismo formato que ver_examen() ──
        'exam': examen,
        'institution': {'name': examen.institution_name or '-'},
        'faculty':     {'name': examen.faculty_name or '-'},
        'career':      {'name': examen.career_name or '-'},
        'subject':     {'name': examen.subject.name if examen.subject else '-'},
        'professor':   {'get_full_name': professor_name},
        'current_date': examen.date_str or '-',
        'exam_type':  _TIPO_EXAMEN_LABELS.get(examen.exam_type or '', examen.exam_type or '-'),
        'exam_mode':  _TIPO_MODALIDAD_LABELS.get(examen.exam_group or '', examen.exam_group or '-'),
        'resolution_time': examen.duration_minutes,
        'modalidad_resolucion': modalidad_list,
        'instructions': examen.instructions or '',
        'questions_texts': questions_texts,
        'outcomes_texts': outcomes_texts,
        'topics_texts': topics_texts,
        # ── extra solo para exportación ──
        'con_respuestas': con_respuestas,
    }


def _b64_to_bytes(data_uri):
    """
    Convierte una cadena data:mime;base64,... a bytes puros.
    Retorna None si el formato no es válido.
    """
    if not data_uri or not data_uri.startswith('data:'):
        return None
    try:
        _header, encoded = data_uri.split(',', 1)
        return base64.b64decode(encoded)
    except Exception:
        return None


def _safe_filename(title, suffix=''):
    """Genera un nombre de archivo seguro a partir del título del examen."""
    safe = ''.join(c if c.isalnum() or c in ' _-' else '_' for c in title)
    safe = safe.strip().replace(' ', '_')[:60]
    return f"{safe}{suffix}"


# ──────────────────────────────────────────────────────────────────────────────
# Exportar a DOCX
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def exportar_examen_docx(request, pk):
    """
    Genera y descarga el examen como archivo .docx (Office Open XML).
    El orden de secciones es idéntico al de base_exam_preview.html:
      encabezado → info (2 cols) → outcomes → instrucciones → temas → preguntas

    Query params:
      ?con_respuestas=1   incluye la clave de respuestas de cada pregunta.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    examen = get_object_or_404(Exam, pk=pk, created_by=request.user)
    con_respuestas = request.GET.get('con_respuestas') == '1'
    ctx = _build_export_context(examen, con_respuestas)

    doc = Document()

    # ── Márgenes de página ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 1. ENCABEZADO INSTITUCIONAL ──
    inst_name = ctx['institution']['name']
    if inst_name and inst_name != '-':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(inst_name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

    fac_name = ctx['faculty']['name']
    car_name = ctx['career']['name']
    if fac_name and fac_name != '-':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{fac_name}  –  {car_name}" if car_name and car_name != '-' else fac_name)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

    subj_name = ctx['subject']['name']
    if subj_name and subj_name != '-':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subj_name)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # ── 2. TABLA DE METADATOS (2 columnas, sin bordes — igual al flex del browser) ──
    info_table = doc.add_table(rows=1, cols=2)
    # Eliminar bordes de la tabla vía XML (python-docx no tiene API directa)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = info_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        border_el = OxmlElement(border_name)
        border_el.set(qn('w:val'), 'none')
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

    left_cell = info_table.rows[0].cells[0]
    right_cell = info_table.rows[0].cells[1]

    def _cell_line(cell, label, value):
        if value and value != '-':
            p = cell.add_paragraph()
            lbl = p.add_run(f"{label}: ")
            lbl.bold = True
            lbl.font.size = Pt(10)
            lbl.font.color.rgb = RGBColor(0, 102, 204)
            val = p.add_run(str(value))
            val.font.size = Pt(10)

    prof = ctx['professor']['get_full_name']
    _cell_line(left_cell,  'Profesor/a', prof)
    _cell_line(left_cell,  'Fecha', ctx['current_date'])
    _cell_line(left_cell,  'Alumno/a', ctx['exam'].alumno or '')
    _cell_line(left_cell,  'Curso/Comisión', ctx['exam'].curso or '')

    _cell_line(right_cell, 'Tipo de examen', ctx['exam_type'])
    _cell_line(right_cell, 'Modalidad', ctx['exam_mode'])
    if ctx['resolution_time']:
        _cell_line(right_cell, 'Duración', f"{ctx['resolution_time']} minutos")
    if ctx['modalidad_resolucion']:
        _cell_line(right_cell, 'Mod. resolución', ', '.join(ctx['modalidad_resolucion']))

    # Limpiar el primer párrafo vacío que agrega docx por defecto en las celdas
    for cell in (left_cell, right_cell):
        if cell.paragraphs and not cell.paragraphs[0].text:
            p = cell.paragraphs[0]._element
            p.getparent().remove(p)

    doc.add_paragraph()

    # ── 3. RESULTADOS DE APRENDIZAJE ──
    if ctx['outcomes_texts']:
        h = doc.add_heading('Resultados de aprendizaje evaluados:', level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for o in ctx['outcomes_texts']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(o))
        doc.add_paragraph()

    # ── 4. INSTRUCCIONES ──
    if ctx['instructions']:
        h = doc.add_heading('Instrucciones generales:', level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph(ctx['instructions'])
        doc.add_paragraph()

    # ── 5. TEMAS ──
    if ctx['topics_texts']:
        h = doc.add_heading('Temas a evaluar:', level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        for t in ctx['topics_texts']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(str(t))
        doc.add_paragraph()

    # ── 6. PREGUNTAS ──
    h = doc.add_heading('Preguntas:', level=2)
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    for q in ctx['questions_texts']:
        q_para = doc.add_paragraph(style='List Number')
        q_run = q_para.add_run(q['text'])
        q_run.bold = True
        q_run.font.size = Pt(11)
        q_run.font.color.rgb = RGBColor(0, 102, 204)  # #0066cc, igual que strong en browser

        if q.get('question_image_b64'):
            img_bytes = _b64_to_bytes(q['question_image_b64'])
            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Cm(12))
                except Exception:
                    pass

        qtype = q.get('type', '')
        if qtype == 'opcion_multiple' and q.get('options'):
            for j, opt in enumerate(q['options']):
                letter = _OPTION_LETTERS[j] if j < len(_OPTION_LETTERS) else str(j + 1)
                opt_p = doc.add_paragraph(style='List Bullet')
                opt_p.add_run(f"{letter}) {opt}")
        elif qtype == 'verdadero_falso':
            p = doc.add_paragraph()
            p.add_run('○ Verdadero          ○ Falso')
        elif qtype == 'desarrollo':
            for _ in range(4):
                blank = doc.add_paragraph('_' * 85)
                blank.runs[0].font.color.rgb = RGBColor(180, 180, 180)
                blank.runs[0].font.size = Pt(9)

        if con_respuestas and (q.get('answer_text') or q.get('answer_image_b64')):
            ans_p = doc.add_paragraph()
            label_run = ans_p.add_run('✓ Respuesta: ')
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(0, 102, 204)
            if q.get('answer_text'):
                ans_p.add_run(q['answer_text'])
            if q.get('answer_image_b64'):
                img_bytes = _b64_to_bytes(q['answer_image_b64'])
                if img_bytes:
                    try:
                        doc.add_picture(io.BytesIO(img_bytes), width=Cm(10))
                    except Exception:
                        pass

    # ── Serializar y responder ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    suffix = '_con_respuestas' if con_respuestas else ''
    filename = _safe_filename(examen.title or 'Examen', suffix) + '.docx'

    response = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ──────────────────────────────────────────────────────────────────────────────
# Exportar a PDF
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def exportar_examen_pdf(request, pk):
    """
    Genera y descarga el examen como archivo .pdf.

    Usa xhtml2pdf (ReportLab) para convertir un template HTML en PDF.
    El template usa exactamente las mismas variables que ver_examen(),
    garantizando que el PDF sea idéntico al resultado de imprimir desde el
    navegador (Ctrl+P).

    Query params:
      ?con_respuestas=1   incluye la clave de respuestas de cada pregunta.
    """
    examen = get_object_or_404(Exam, pk=pk, created_by=request.user)
    con_respuestas = request.GET.get('con_respuestas') == '1'
    ctx = _build_export_context(examen, con_respuestas)

    try:
        from xhtml2pdf import pisa
    except Exception as _xhtml_err:
        return HttpResponse(
            f'Error al cargar xhtml2pdf: {_xhtml_err}\n\n'
            'Asegúrese de que xhtml2pdf esté instalado en el mismo entorno virtual '
            'que usa Django y reinicie el servidor.',
            status=503,
            content_type='text/plain; charset=utf-8',
        )

    html_string = render_to_string('material/exams/exam_export_pdf.html', ctx, request=request)

    buf = io.BytesIO()
    pisa_status = pisa.pisaDocument(io.BytesIO(html_string.encode('utf-8')), buf)

    if pisa_status.err:
        return HttpResponse(
            'Error al generar el PDF. Por favor, use la función de impresión del navegador.',
            status=500,
        )

    buf.seek(0)
    suffix = '_con_respuestas' if con_respuestas else ''
    filename = _safe_filename(examen.title or 'Examen', suffix) + '.pdf'

    response = HttpResponse(buf.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
