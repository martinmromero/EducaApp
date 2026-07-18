import io
import os
import base64
from PIL import Image as PILImage


def render_exam_payload_to_docx(payload, formato):
    """Renderiza payload neutral de examen a DOCX y retorna bytes."""
    from docx import Document

    doc = Document()
    _apply_margins(doc, formato)
    _append_payload(doc, payload, formato)
    _enforce_font_name_in_document(doc, getattr(formato, 'fuente', 'Arial') or 'Arial')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def render_exam_batch_payloads_to_docx(exam_documents):
    """Renderiza varios payloads de examen a un unico DOCX y retorna bytes."""
    from docx import Document

    doc = Document()
    if not exam_documents:
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out.read()

    _apply_margins(doc, exam_documents[0]['formato'])

    for idx, item in enumerate(exam_documents):
        formato = item['formato']
        payload = item['payload']

        if idx > 0:
            doc.add_page_break()

        _append_payload(doc, payload, formato)

    _enforce_font_name_in_document(doc, getattr(exam_documents[0]['formato'], 'fuente', 'Arial') or 'Arial')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _apply_margins(doc, formato):
    from docx.shared import Cm

    for section in doc.sections:
        section.top_margin = Cm(float(getattr(formato, 'margen_superior_cm', 2.0) or 2.0))
        section.bottom_margin = Cm(float(getattr(formato, 'margen_inferior_cm', 2.0) or 2.0))
        section.left_margin = Cm(float(getattr(formato, 'margen_izquierdo_cm', 2.5) or 2.5))
        section.right_margin = Cm(float(getattr(formato, 'margen_derecho_cm', 2.0) or 2.0))


def _set_run_font(run, *, font_name, size_pt=None, color_rgb=None, bold=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if bold is not None:
        run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)

    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:cs'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)


def _add_line(doc, text, *, base_size, text_rgb, font_name, bold=False, size=None, color=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        font_name=font_name,
        size_pt=size or base_size,
        color_rgb=(color or text_rgb),
        bold=bold,
    )
    return paragraph


def _append_payload(doc, payload, formato):
    base_size = int(getattr(formato, 'tamano_fuente', 11) or 11)
    font_name = getattr(formato, 'fuente', 'Arial') or 'Arial'
    title_rgb = _hex_to_rgb(getattr(formato, 'color_titulo', '') or '#111111')
    text_rgb = _hex_to_rgb(getattr(formato, 'color_texto', '') or '#111111')

    def add_line(text, bold=False, size=None, color=None):
        return _add_line(
            doc,
            text,
            base_size=base_size,
            text_rgb=text_rgb,
            font_name=font_name,
            bold=bold,
            size=size,
            color=color,
        )

    for block in payload:
        tipo = block.get('tipo')

        if tipo in {'letterhead', 'encabezado'}:
            _append_letterhead_table(doc, block, base_size, title_rgb, text_rgb, font_name)

        elif tipo == 'datos_alumno':
            _append_student_data_table(doc, block, base_size, title_rgb, text_rgb, font_name)

        elif tipo == 'titulo' and block.get('texto'):
            p = add_line(block['texto'], bold=True, size=base_size + 4, color=title_rgb)
            p.alignment = 1

        elif tipo in {'instrucciones', 'instrucciones_generales'} and block.get('texto'):
            add_line('Instrucciones generales', bold=True, size=base_size + 1, color=title_rgb)
            add_line(block['texto'])

        elif tipo in {'lista_outcomes', 'resultados_aprendizaje'} and block.get('items'):
            add_line('Resultados de aprendizaje evaluados', bold=True, size=base_size + 1, color=title_rgb)
            for item in block['items']:
                add_line(f"- {item}")

        elif tipo == 'requisitos_aprobar' and block.get('texto'):
            add_line('Requisitos para aprobar', bold=True, size=base_size + 1, color=title_rgb)
            add_line(block['texto'])

        elif tipo == 'tiempo' and block.get('texto'):
            add_line(f"Tiempo: {block['texto']}", bold=False)

        elif tipo == 'lista_temas' and block.get('items'):
            add_line('Temas a evaluar', bold=True, size=base_size + 1, color=title_rgb)
            for item in block['items']:
                add_line(f"- {item}")

        elif tipo == 'pregunta':
            add_line(f"{block.get('numero', '')}. {block.get('texto', '')}", bold=True)
            subtype = block.get('subtipo')
            if subtype == 'opcion_multiple':
                for idx, opt in enumerate(block.get('opciones', []), start=1):
                    add_line(f"   {idx}) {opt}")
            elif subtype == 'verdadero_falso':
                add_line('   ( ) Verdadero      ( ) Falso')
            elif subtype in ('completar_blank', 'desarrollo'):
                add_line('   ___________________________________________')
                add_line('   ___________________________________________')
            if block.get('respuesta'):
                add_line(f"Respuesta: {block['respuesta']}", bold=True, color=title_rgb)

        elif tipo == 'tabla_rubrica' and block.get('filas'):
            add_line(block.get('titulo', 'Rubrica'), bold=True, size=base_size + 1, color=title_rgb)
            headers = ['Criterio'] + list(block.get('columnas', []))
            table = doc.add_table(rows=1, cols=len(headers))
            _set_table_fixed_layout(table)
            _apply_table_borders(table)
            for i, h in enumerate(headers):
                _set_cell_text(
                    table.rows[0].cells[i],
                    h,
                    font_name=font_name,
                    size_pt=base_size,
                    color_rgb=title_rgb,
                    bold=True,
                )
            for row in block['filas']:
                cells = table.add_row().cells
                _set_cell_text(cells[0], row.get('criterio', ''), font_name=font_name, size_pt=base_size, color_rgb=text_rgb)
                for i, value in enumerate(row.get('celdas', []), start=1):
                    if i < len(cells):
                        _set_cell_text(cells[i], value, font_name=font_name, size_pt=base_size, color_rgb=text_rgb)


def _hex_to_rgb(hex_color):
    value = (hex_color or '').strip()
    if not (len(value) == 7 and value.startswith('#')):
        value = '#111111'
    return int(value[1:3], 16), int(value[3:5], 16), int(value[5:7], 16)


def _decode_data_uri(data_uri):
    if not data_uri or not str(data_uri).startswith('data:'):
        return None
    try:
        _, encoded = str(data_uri).split(',', 1)
        return base64.b64decode(encoded)
    except Exception:
        return None


def _resolve_logo_bytes_or_path(block):
    logo_b64 = block.get('logo_b64') or ''
    logo_url = block.get('logo_url') or ''
    logo_path = block.get('logo_path') or ''

    img_bytes = _decode_data_uri(logo_b64) or _decode_data_uri(logo_url)
    if img_bytes:
        return io.BytesIO(img_bytes), None

    for candidate in [logo_path, logo_url]:
        if candidate and os.path.exists(candidate):
            return None, candidate
    return None, None


def _apply_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge = OxmlElement(f'w:{name}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '8')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), '444444')
        borders.append(edge)
    tbl_pr.append(borders)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for name in ['top', 'left', 'bottom', 'right']:
                edge = OxmlElement(f'w:{name}')
                edge.set(qn('w:val'), 'single')
                edge.set(qn('w:sz'), '8')
                edge.set(qn('w:space'), '0')
                edge.set(qn('w:color'), '444444')
                tc_borders.append(edge)
            tc_pr.append(tc_borders)


def _set_table_fixed_layout(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)


def _set_column_width(table, col_idx, width_cm):
    from docx.shared import Cm

    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _clear_cell(cell):
    cell.text = ''
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph('')
    paragraph.clear()
    return paragraph


def _set_cell_text(cell, text, *, font_name, size_pt, color_rgb, bold=False, alignment=None):
    paragraph = _clear_cell(cell)
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(str(text or ''))
    _set_run_font(
        run,
        font_name=font_name,
        size_pt=size_pt,
        color_rgb=color_rgb,
        bold=bold,
    )
    return paragraph, run


def _append_run(paragraph, text, *, font_name, size_pt, color_rgb, bold=False):
    run = paragraph.add_run(str(text or ''))
    _set_run_font(
        run,
        font_name=font_name,
        size_pt=size_pt,
        color_rgb=color_rgb,
        bold=bold,
    )
    return run


def _enforce_font_name_in_document(doc, font_name):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name=font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, font_name=font_name)


def _append_letterhead_table(doc, block, base_size, title_rgb, text_rgb, font_name):
    from docx.shared import Cm

    table = doc.add_table(rows=2, cols=3)
    _set_table_fixed_layout(table)
    _set_column_width(table, 0, 2.1)
    _set_column_width(table, 2, 2.3)
    _apply_table_borders(table)

    institution = (block.get('institucion') or '-').upper()
    career = block.get('carrera') or '-'
    subject = block.get('materia') or '-'
    professor = block.get('profesor') or '-'
    exam_type = block.get('tipo_examen') or 'Examen'
    year = str(block.get('anio') or '-')

    logo_cell = table.rows[0].cells[0].merge(table.rows[1].cells[0])
    year_cell = table.rows[0].cells[2].merge(table.rows[1].cells[2])

    logo_stream, logo_path = _resolve_logo_bytes_or_path(block)
    logo_paragraph = _clear_cell(logo_cell)
    logo_paragraph.alignment = 1
    if logo_stream or logo_path:
        run = logo_paragraph.add_run()
        try:
            width_cm, height_cm, constrain_by = _fit_logo_size_for_docx(logo_stream=logo_stream, logo_path=logo_path)
            if constrain_by == 'width':
                run.add_picture(logo_stream or logo_path, width=Cm(width_cm))
            else:
                run.add_picture(logo_stream or logo_path, height=Cm(height_cm))
        except Exception:
            pass
        _set_run_font(run, font_name=font_name)

    center_top = table.rows[0].cells[1].paragraphs[0]
    center_top.clear()
    center_top.alignment = 1
    _append_run(
        center_top,
        institution,
        font_name=font_name,
        size_pt=base_size + 1,
        color_rgb=title_rgb,
        bold=True,
    )

    right_top = _clear_cell(year_cell)
    right_top.alignment = 1
    _append_run(
        right_top,
        'Año',
        font_name=font_name,
        size_pt=base_size - 1 if base_size > 9 else base_size,
        color_rgb=text_rgb,
        bold=True,
    )
    right_top.add_run('\n')
    _append_run(
        right_top,
        year,
        font_name=font_name,
        size_pt=base_size,
        color_rgb=text_rgb,
        bold=False,
    )

    meta_cell = table.rows[1].cells[1]
    p_meta = _clear_cell(meta_cell)
    p_meta.alignment = 0
    _append_run(p_meta, 'Carrera: ', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _append_run(p_meta, career, font_name=font_name, size_pt=base_size, color_rgb=text_rgb)
    _append_run(p_meta, '    Profesor: ', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _append_run(p_meta, professor, font_name=font_name, size_pt=base_size, color_rgb=text_rgb)

    p_meta = meta_cell.add_paragraph()
    _append_run(p_meta, 'Materia: ', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _append_run(p_meta, subject, font_name=font_name, size_pt=base_size, color_rgb=text_rgb)
    _append_run(p_meta, '    ', font_name=font_name, size_pt=base_size, color_rgb=text_rgb)
    _append_run(p_meta, exam_type, font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)


def _append_student_data_table(doc, block, base_size, title_rgb, text_rgb, font_name):
    table = doc.add_table(rows=2, cols=4)
    _set_table_fixed_layout(table)
    _set_column_width(table, 0, 2.4)
    _set_column_width(table, 2, 2.4)
    _apply_table_borders(table)

    _set_cell_text(table.rows[0].cells[0], 'Nombre', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _set_cell_text(table.rows[0].cells[1], '', font_name=font_name, size_pt=base_size, color_rgb=text_rgb)
    _set_cell_text(table.rows[0].cells[2], 'Apellido', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _set_cell_text(table.rows[0].cells[3], '', font_name=font_name, size_pt=base_size, color_rgb=text_rgb)

    _set_cell_text(table.rows[1].cells[0], 'Fecha', font_name=font_name, size_pt=base_size, color_rgb=text_rgb, bold=True)
    _set_cell_text(table.rows[1].cells[1], block.get('fecha') or '', font_name=font_name, size_pt=base_size, color_rgb=text_rgb)

    merged = table.rows[1].cells[2].merge(table.rows[1].cells[3])
    _set_cell_text(
        merged,
        block.get('tipo_examen_mayusculas') or (block.get('tipo_examen') or 'EXAMEN').upper(),
        font_name=font_name,
        size_pt=base_size,
        color_rgb=title_rgb,
        bold=True,
        alignment=1,
    )


def _fit_logo_size_for_docx(logo_stream=None, logo_path=None):
    max_width_cm = 1.6
    max_height_cm = 1.2

    try:
        if logo_stream is not None:
            logo_stream.seek(0)
            image = PILImage.open(logo_stream)
        else:
            image = PILImage.open(logo_path)

        original_width, original_height = image.size
        image.close()
        if not original_width or not original_height:
            return max_width_cm, max_height_cm, 'width'

        width_scale = max_width_cm / float(original_width)
        height_scale = max_height_cm / float(original_height)
        scale = min(width_scale, height_scale, 1.0)

        width_cm = original_width * scale
        height_cm = original_height * scale
        if logo_stream is not None:
            logo_stream.seek(0)

        if width_scale <= height_scale:
            return width_cm, height_cm, 'width'
        return width_cm, height_cm, 'height'
    except Exception:
        if logo_stream is not None:
            logo_stream.seek(0)
        return max_width_cm, max_height_cm, 'width'
